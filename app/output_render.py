"""输出复刻层（场景无关）。

红线 #3：产出默认复刻**历史结果文件的格式**——历史结果是 Excel 就产出 Excel，
是 CSV/MD/Word 就产出同格式，列结构与历史结果一致。用户也可在对话中要求改用其他展示
（如临时在中间区域看表、导出 CSV/MD），由调用方覆盖 `fmt`。

本模块只负责「把一个结果 DataFrame 按指定格式写成文件」，不含任何业务逻辑、不写死字段名。
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

import pandas as pd

# 历史结果文件扩展名 → 输出格式
_EXT_TO_FMT = {
    ".xlsx": "xlsx", ".xls": "xlsx",
    ".csv": "csv", ".tsv": "tsv",
    ".md": "md", ".markdown": "md",
    ".docx": "docx", ".doc": "docx",
    ".json": "json",
}

SUPPORTED_FORMATS = {"xlsx", "csv", "tsv", "md", "docx", "json"}


def infer_format(file_path: str | Path) -> str:
    """据历史结果文件扩展名推断产出格式；未知则退化为 csv。"""
    ext = Path(file_path).suffix.lower()
    return _EXT_TO_FMT.get(ext, "csv")


def _safe_stem(name: str) -> str:
    keep = []
    for ch in str(name):
        keep.append(ch if (ch.isalnum() or ch in "._-（）()[]中文" or ord(ch) > 0x2E80) else "_")
    stem = "".join(keep).strip("._") or "output"
    return stem[:80]


def render(
    df: pd.DataFrame,
    fmt: str,
    out_dir: str | Path,
    base_name: str,
    columns: Optional[list[str]] = None,
) -> Path:
    """把结果 DataFrame 按 `fmt` 写到 `out_dir/base_name.<ext>`，返回文件路径。

    - columns 给定时，先对齐到该列契约（缺列补空、定序、丢弃多余列）。
    - 不支持/失败的格式安全回退为 csv，保证总能产出文件。
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    if df is None:
        df = pd.DataFrame()
    if columns:
        df = df.reindex(columns=columns)

    fmt = (fmt or "csv").lower()
    stem = _safe_stem(base_name)
    try:
        return _RENDERERS.get(fmt, _to_csv)(df, out_dir, stem)
    except Exception:  # noqa: BLE001  任何渲染失败都回退 csv，不让产出落空
        return _to_csv(df, out_dir, stem)


# ---------------------------------------------------------------------------
def _to_csv(df: pd.DataFrame, out_dir: Path, stem: str) -> Path:
    path = out_dir / f"{stem}.csv"
    df.to_csv(path, index=False, encoding="utf-8-sig")
    return path


def _to_tsv(df: pd.DataFrame, out_dir: Path, stem: str) -> Path:
    path = out_dir / f"{stem}.tsv"
    df.to_csv(path, index=False, sep="\t", encoding="utf-8-sig")
    return path


def _to_xlsx(df: pd.DataFrame, out_dir: Path, stem: str) -> Path:
    path = out_dir / f"{stem}.xlsx"
    df.to_excel(path, index=False)
    return path


def _to_md(df: pd.DataFrame, out_dir: Path, stem: str) -> Path:
    path = out_dir / f"{stem}.md"
    try:
        text = df.to_markdown(index=False)
    except Exception:  # noqa: BLE001  缺 tabulate 时手工拼一张 Markdown 表
        text = _manual_markdown(df)
    path.write_text(f"# {stem}\n\n共 {len(df)} 行。\n\n{text}\n", encoding="utf-8")
    return path


def _to_json(df: pd.DataFrame, out_dir: Path, stem: str) -> Path:
    path = out_dir / f"{stem}.json"
    df.to_json(path, orient="records", force_ascii=False, indent=2)
    return path


def _to_docx(df: pd.DataFrame, out_dir: Path, stem: str) -> Path:
    """写 Word 文档（需 python-docx）；不可用时回退为 Markdown。"""
    try:
        from docx import Document
    except Exception:  # noqa: BLE001
        return _to_md(df, out_dir, stem)
    doc = Document()
    doc.add_heading(stem, level=1)
    doc.add_paragraph(f"共 {len(df)} 行。")
    cols = [str(c) for c in df.columns]
    table = doc.add_table(rows=1, cols=max(len(cols), 1))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = c
    for _, row in df.head(5000).iterrows():  # Word 表过大易卡，限量
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = "" if pd.isna(row[c]) else str(row[c])
    path = out_dir / f"{stem}.docx"
    doc.save(path)
    return path


def _manual_markdown(df: pd.DataFrame) -> str:
    cols = [str(c) for c in df.columns] or ["(空)"]
    head = "| " + " | ".join(cols) + " |"
    sep = "| " + " | ".join("---" for _ in cols) + " |"
    body = []
    for _, row in df.head(1000).iterrows():
        body.append("| " + " | ".join(
            "" if pd.isna(row[c]) else str(row[c]) for c in df.columns
        ) + " |")
    return "\n".join([head, sep, *body])


_RENDERERS = {
    "csv": _to_csv,
    "tsv": _to_tsv,
    "xlsx": _to_xlsx,
    "md": _to_md,
    "json": _to_json,
    "docx": _to_docx,
}
